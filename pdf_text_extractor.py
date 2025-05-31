"""
PDF Text Extractor with OCR
---------------------------------------
This script allows you to extract text from PDF documents using either direct text extraction or Optical Character Recognition (OCR). 
It's especially useful for scanned PDFs or image-based documents that don't contain selectable text.
by Denis (BeforeMyCompileFails) 2025
"""

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import os
import sys
import io

def setup_tesseract():
    """Setup Tesseract OCR path for Windows"""
    # Common Tesseract installation paths on Windows
    possible_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\%USERNAME%\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
    ]
    
    for path in possible_paths:
        expanded_path = os.path.expandvars(path)
        if os.path.exists(expanded_path):
            pytesseract.pytesseract.tesseract_cmd = expanded_path
            return True
    
    # If not found, try using it from PATH
    try:
        pytesseract.get_tesseract_version()
        return True
    except:
        return False

def extract_text_with_ocr(pdf_path):
    """Extract text from PDF using OCR on page images"""
    try:
        # Open PDF
        doc = fitz.open(pdf_path)
        full_text = ""
        
        print(f"Processing {len(doc)} pages with OCR...")
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Convert page to image
            mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR accuracy
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            try:
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text.strip():
                    full_text += f"\n--- Page {page_num + 1} ---\n"
                    full_text += page_text + "\n"
                    print(f"✓ Extracted text from page {page_num + 1}")
                else:
                    print(f"⚠ Page {page_num + 1} contains no readable text")
            except Exception as e:
                print(f"✗ Error processing page {page_num + 1}: {e}")
        
        doc.close()
        return full_text
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return None

def extract_text_simple(pdf_path):
    """Try simple text extraction first (faster)"""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page_text + "\n"
        
        doc.close()
        return text if text.strip() else None
        
    except Exception as e:
        print(f"Error in simple extraction: {e}")
        return None

def save_to_docx(text, output_path):
    """Save extracted text to Word document"""
    try:
        doc = Document()
        doc.add_heading('Extracted PDF Text', 0)
        
        # Split text into paragraphs and add to document
        paragraphs = text.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        doc.save(output_path)
        print(f"✓ Text saved to: {output_path}")
        return True
    except Exception as e:
        print(f"✗ Error saving to Word document: {e}")
        return False

def save_to_txt(text, output_path):
    """Save extracted text to plain text file"""
    try:
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(text)
        print(f"✓ Text saved to: {output_path}")
        return True
    except Exception as e:
        print(f"✗ Error saving to text file: {e}")
        return False

def main():
    print("Advanced PDF Text Extractor with OCR")
    print("=" * 50)
    
    # Check if Tesseract is available
    if not setup_tesseract():
        print("⚠ Warning: Tesseract OCR not found!")
        print("Please install Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki")
        print("The script will try simple extraction only.\n")
    
    # Get PDF file path
    if len(sys.argv) > 1:
        pdf_path = sys.argv[1]
    else:
        pdf_path = input("Enter the path to your PDF file: ").strip().strip('"')
    
    # Check if file exists
    if not os.path.exists(pdf_path):
        print(f"✗ Error: File '{pdf_path}' not found!")
        return
    
    print(f"📄 Processing: {pdf_path}")
    
    # Try simple extraction first
    print("\n🔍 Attempting simple text extraction...")
    extracted_text = extract_text_simple(pdf_path)
    
    if not extracted_text or not extracted_text.strip():
        print("⚠ Simple extraction failed or returned no text.")
        print("🔍 Attempting OCR extraction...")
        
        if not setup_tesseract():
            print("✗ Cannot perform OCR without Tesseract. Please install it first.")
            return
        
        extracted_text = extract_text_with_ocr(pdf_path)
    else:
        print("✓ Simple extraction successful!")
    
    if not extracted_text or not extracted_text.strip():
        print("✗ Failed to extract any text from PDF!")
        return
    
    # Generate output filenames
    base_name = os.path.splitext(pdf_path)[0]
    docx_output = f"{base_name}_extracted.docx"
    txt_output = f"{base_name}_extracted.txt"
    
    # Ask user for output format preference
    print(f"\n📊 Extracted {len(extracted_text)} characters")
    print("\nChoose output format:")
    print("1. Word document (.docx)")
    print("2. Plain text file (.txt)")
    print("3. Both")
    
    choice = input("Enter your choice (1-3): ").strip()
    
    success = False
    if choice in ['1', '3']:
        success = save_to_docx(extracted_text, docx_output) or success
    if choice in ['2', '3']:
        success = save_to_txt(extracted_text, txt_output) or success
    
    if success:
        print(f"\n🎉 Extraction completed successfully!")
        print(f"📈 Total characters extracted: {len(extracted_text)}")
    else:
        print("✗ Failed to save extracted text!")

if __name__ == "__main__":
    main()